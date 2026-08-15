import os
import shutil
import json
import re
import tempfile
import zipfile
import concurrent.futures
import streamlit as st
import time
from google import genai
from google.genai import types
from PIL import Image
from pypdf import PdfReader
import fitz  # PyMuPDF
from docx import Document
import pypandoc

def get_api_key():
    """Tenta di caricare l'API key di Gemini da variabili d'ambiente o st.secrets."""
    # 1. Prova da variabile d'ambiente
    api_key = os.environ.get("GEMINI_API_KEY")
    if api_key:
        return api_key.strip()
        
    # 2. Prova da st.secrets
    try:
        if "GEMINI_API_KEY" in st.secrets:
            return st.secrets["GEMINI_API_KEY"].strip()
    except:
        pass
            
    return None


# Configurazione della pagina
st.set_page_config(
    page_title="Risolutore Matematica & Fisica",
    page_icon="📐",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Custom per applicare un look pastello premium
st.markdown(
    """
    <style>
    /* Sfondo principale e font */
    .stApp {
        background-color: #F8F9FA;
        color: #2D3748;
        font-family: 'Inter', 'Outfit', 'Helvetica Neue', sans-serif;
    }
    
    /* Titoli principali */
    h1, h2, h3 {
        color: #2D3748;
        font-weight: 700;
    }
    
    /* Barra laterale pastello */
    section[data-testid="stSidebar"] {
        background-color: #E3F2FD !important;
        border-right: 1px solid #CFD8DC;
    }
    
    /* Contenitori e schede */
    .stAlert {
        border-radius: 8px;
        border: none;
    }
    
    /* Bottone Call to Action principale (Ottanio desaturato) */
    div.stButton > button[kind="primary"],
    div.stButton > button[data-testid="baseButton-primary"],
    div.stDownloadButton > button[kind="primary"] {
        background-color: #2A9D8F !important;
        color: white !important;
        border: none !important;
        border-radius: 8px;
        padding: 0.6rem 1.8rem;
        font-weight: 600;
        font-size: 1rem;
        transition: all 0.3s ease;
        box-shadow: 0px 4px 6px rgba(42, 157, 143, 0.2);
    }
    
    div.stButton > button[kind="primary"]:hover,
    div.stButton > button[data-testid="baseButton-primary"]:hover,
    div.stDownloadButton > button[kind="primary"]:hover {
        background-color: #218376 !important;
        color: white !important;
        box-shadow: 0px 6px 12px rgba(33, 131, 118, 0.3);
        transform: translateY(-2px);
    }
    
    div.stButton > button:not([kind="primary"]):not([data-testid="baseButton-primary"]) {
        border-radius: 8px;
        font-weight: 600;
    }
    
    /* File Uploader styling */
    section[data-testid="stFileUploadDropzone"] {
        border: 2px dashed #2A9D8F;
        background-color: #FFFFFF;
        border-radius: 8px;
    }
    
    /* Nascondi elementi nativi di Streamlit (Menu tre puntini, Fork, GitHub, Footer e Badge Cloud) */
    #MainMenu {visibility: hidden !important; display: none !important;}
    header[data-testid="stHeader"] {visibility: hidden !important; display: none !important;}
    footer {visibility: hidden !important; display: none !important;}
    div[data-testid="stToolbar"] {visibility: hidden !important; display: none !important;}
    div[data-testid="stDecoration"] {visibility: hidden !important; display: none !important;}
    div[data-testid="stStatusWidget"] {visibility: hidden !important; display: none !important;}
    [data-testid="manage-app-button"] {visibility: hidden !important; display: none !important;}
    [data-testid="stAppDeployButton"] {visibility: hidden !important; display: none !important;}
    
    /* Badge 'Hosted with Streamlit', avatar e pulsanti galleggianti in basso a destra */
    div[class*="viewerBadge"],
    div[class*="ViewerBadge"],
    a[class*="viewerBadge"],
    a[class*="ViewerBadge"],
    .viewerBadge_container__1QSob,
    .viewerBadge_link__1S137,
    .viewerBadge_text__1JaDK,
    div[class*="stDeployButton"],
    div[class*="floating-button"],
    div[class*="FloatingButton"],
    div[class*="bottom-bar"],
    div[class*="BottomBar"],
    div[class*="HostBadge"],
    div[class*="hostBadge"],
    div[class*="Avatar"],
    div[class*="avatar"] {
        display: none !important;
        visibility: hidden !important;
        opacity: 0 !important;
        pointer-events: none !important;
        height: 0 !important;
        width: 0 !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Funzioni di estrazione per il Modulo 1

def clean_temp_images(output_dir):
    """Rimuove i file di immagini temporanee accumulati nella cartella."""
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)

def extract_images_from_docx(docx_path, output_dir, source_prefix=""):
    """Estrae tutte le immagini incorporate in un file DOCX."""
    doc = Document(docx_path)
    extracted_images = []
    prefix = f"{source_prefix}_" if source_prefix else ""
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                image_name = os.path.basename(rel.target_ref)
                # Crea un nome univoco per evitare sovrascritture
                unique_name = f"{prefix}docx_img_{rel_id}_{image_name}"
                dest_path = os.path.join(output_dir, unique_name)
                with open(dest_path, "wb") as f:
                    f.write(image_data)
                extracted_images.append(unique_name)
            except Exception as e:
                st.warning(f"Impossibile estrarre un'immagine da DOCX: {e}")
    return extracted_images

def extract_images_from_pdf(pdf_path, output_dir, source_prefix=""):
    """Estrae le immagini raster incorporate nelle pagine del PDF tramite pypdf."""
    reader = PdfReader(pdf_path)
    extracted_images = []
    prefix = f"{source_prefix}_" if source_prefix else ""
    for page_idx, page in enumerate(reader.pages):
        for img_idx, img_file_object in enumerate(page.images):
            try:
                unique_name = f"{prefix}pdf_p{page_idx}_img{img_idx}_{img_file_object.name}"
                dest_path = os.path.join(output_dir, unique_name)
                with open(dest_path, "wb") as fp:
                    fp.write(img_file_object.data)
                extracted_images.append(unique_name)
            except Exception as e:
                st.warning(f"Impossibile estrarre un'immagine dal PDF (p. {page_idx}): {e}")
    return extracted_images

def convert_pdf_pages_to_images(pdf_path, output_dir):
    """Rasterizza ogni pagina del PDF in un'immagine PNG (per PDF scansionati o scritti a mano)."""
    doc = fitz.open(pdf_path)
    if len(doc) > 20:
        st.error(f"Il PDF supera il limite di 20 pagine.")
        st.stop()
    page_images = []
    for page_num in range(len(doc)):
        try:
            page = doc.load_page(page_num)
            pix = page.get_pixmap(dpi=150)
            if pix.width * pix.height > 20_000_000:
                st.error(f"La pagina {page_num} supera il limite di 20 megapixel.")
                st.stop()
            img_name = f"page_raster_{page_num}.png"
            dest_path = os.path.join(output_dir, img_name)
            pix.save(dest_path)
            page_images.append(dest_path)
        except Exception as e:
            st.error(f"Errore durante la conversione in immagine della pagina {page_num}: {e}")
    return page_images

def process_input_file(uploaded_file, temp_dir, output_dir):
    """Salva ed elabora il file caricato estraendone testo e immagini."""
    if uploaded_file.size > 15 * 1024 * 1024:
        st.error(f"Il file '{uploaded_file.name}' supera il limite di 15 MB.")
        st.stop()
        
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    base_name = os.path.splitext(uploaded_file.name)[0]
    safe_prefix = re.sub(r'[^a-zA-Z0-9]', '_', base_name)
    
    # Crea un percorso locale temporaneo per il file caricato
    temp_file_path = os.path.join(temp_dir, uploaded_file.name)
    with open(temp_file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
        
    text_content = ""
    images_for_gemini = []  # Immagini dell'intero documento da mostrare a Gemini
    extracted_images = []   # Immagini/grafici interni da mappare
    
    if file_ext == ".txt":
        with open(temp_file_path, "r", encoding="utf-8", errors="ignore") as f:
            text_content = f.read()
            
    elif file_ext == ".docx":
        doc = Document(temp_file_path)
        # Estrai testo da paragrafi
        text_content = "\n".join([p.text for p in doc.paragraphs])
        # Estrai testo da tabelle
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                text_content += f"\n{row_text}"
        # Estrai immagini incorporate
        extracted_images = extract_images_from_docx(temp_file_path, output_dir, source_prefix=safe_prefix)
        
    elif file_ext == ".pdf":
        # Prova prima ad estrarre il testo digitale
        reader = PdfReader(temp_file_path)
        if len(reader.pages) > 20:
            st.error(f"Il file '{uploaded_file.name}' supera il limite di 20 pagine.")
            st.stop()
        extracted_text_list = []
        for page in reader.pages:
            extracted_text_list.append(page.extract_text() or "")
        text_content = "\n".join(extracted_text_list)
        
        # Estrai immagini incorporate (grafici o schemi)
        extracted_images = extract_images_from_pdf(temp_file_path, output_dir, source_prefix=safe_prefix)
        
        # Rasterizzazione multimodale ad alta definizione (sempre attiva per PDF):
        # Genera immagini ad alta definizione (150 DPI) per consentire a Gemini di visualizzare
        # direttamente le formule matematiche e i grafici con fedeltà assoluta al 100%.
        images_for_gemini = convert_pdf_pages_to_images(temp_file_path, output_dir)
            
    elif file_ext in [".png", ".jpg", ".jpeg"]:
        try:
            img = Image.open(temp_file_path)
            if img.width * img.height > 20_000_000:
                st.error(f"L'immagine '{uploaded_file.name}' supera il limite di 20 megapixel.")
                st.stop()
        except:
            pass
        images_for_gemini = [temp_file_path]
        
    return {
        "filename": uploaded_file.name,
        "text": text_content,
        "images_for_gemini": images_for_gemini,
        "extracted_images": extracted_images
    }


# Modulo 2: Integrazione Gemini per parsing e risoluzione

def identify_exercises_with_gemini(api_key, text_content, images_for_gemini, all_extracted_images):
    """Invia il testo e le immagini a Gemini per dividere il documento in singoli esercizi."""
    client = genai.Client(api_key=api_key)
    
    prompt = f"""Analizza il documento fornito (può contenere testo lineare, tabelle o immagini di compiti scritti a mano) e identifica tutti i singoli problemi o esercizi di matematica e fisica presenti.

CORREGGI E RICOSTRUISCI IL TESTO: Molti caratteri matematici potrebbero essere corrotti dall'estrazione del testo (es. simboli strani, quadratini '█', entità HTML corrotte, pedici errati). Ricostruisci il testo in italiano corretto, rendendo leggibili tutte le formule e le variabili matematiche. Formatta tutte le variabili e le formule matematiche nel testo dell'esercizio usando la sintassi LaTeX inline (es. $m_1 = 0,200 \\text{{ kg}}$, $T_{{Cu}} = 150 \\text{{ °C}}$, $T_{{eq}}$, ecc.). Non aggiungere frasi introduttive o di commento all'infuori del testo pulito del problema.

Fornisci in output un array JSON contenente gli esercizi individuati, seguendo rigorosamente questa struttura:
[
  {{
    "id": 1,
    "testo": "Testo corretto, pulito e formattato in LaTeX dell'esercizio 1 (tradotto in italiano se l'originale è in un'altra lingua)",
    "immagini_correlate": []
  }},
  {{
    "id": 2,
    "testo": "Testo corretto, pulito e formattato in LaTeX dell'esercizio 2",
    "immagini_correlate": ["nome_immagine_estratta.png"]
  }}
]

Regole per il parsing:
1. Identifica ciascun esercizio separatamente. Se ci sono più quesiti correlati all'interno di un unico problema (es. punto a, b, c), tienili uniti all'interno dello stesso esercizio (stesso id).
2. MAPPATURA IMMAGINI E GRAFICI (CRITICO): Se un esercizio contiene o fa riferimento a una figura, grafico cartesiano, schema o tabella presente nel documento, DEVI inserire il nome esatto del relativo file immagine nell'array 'immagini_correlate' dell'esercizio corrispondente.
   Elenco immagini estratte dal documento:
   {json.dumps(all_extracted_images)}
3. Restituisci esclusivamente il codice JSON valido, senza blocchi di markdown (come ```json) o testo aggiuntivo.
"""
    contents = [prompt]
    if text_content:
        contents.append(f"Testo digitale estratto:\n{text_content}")
        
    # Carica le immagini delle pagine rasterizzate per Gemini
    for img_path in images_for_gemini:
        if os.path.exists(img_path):
            img = Image.open(img_path)
            contents.append(img)
            
    primary_model = "gemini-3.7-flash"
    fallback_models = ["gemini-3.5-flash", "gemini-2.5-pro", "gemini-3.1-flash-lite"]
    
    max_retries = 3
    retry_delay = 20  # secondi
    
    response = None
    
    # 1. Tentativo con il modello primario e retry
    for attempt in range(max_retries):
        try:
            response = client.models.generate_content(
                model=primary_model,
                contents=contents,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json"
                )
            )
            data = json.loads(response.text.strip())
            return data
        except Exception as e:
            err_msg = str(e)
            # Se la risorsa è esaurita (429) o il server è sovraccarico (503), attendi e riprova
            if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg or "503" in err_msg or "UNAVAILABLE" in err_msg:
                if attempt < max_retries - 1:
                    st.warning(f"⚠️ Modello {primary_model} sovraccarico o limite quota superato. Attesa di {retry_delay} secondi prima del tentativo {attempt + 2}/{max_retries}...")
                    time.sleep(retry_delay)
                    continue
            
            # Se i tentativi falliscono per quota/sovraccarico/non trovato, andiamo in fallback
            if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg or "503" in err_msg or "UNAVAILABLE" in err_msg or "404" in err_msg:
                st.warning(f"⚠️ Modello primario ({primary_model}) non disponibile. Tento il recupero con i modelli di backup in cascata...")
                
                for fallback_model in fallback_models:
                    st.info(f"🔄 Tentativo con modello di backup: {fallback_model}...")
                    for fb_attempt in range(max_retries):
                        try:
                            response = client.models.generate_content(
                                model=fallback_model,
                                contents=contents,
                                config=types.GenerateContentConfig(
                                    response_mime_type="application/json"
                                )
                            )
                            data = json.loads(response.text.strip())
                            return data
                        except Exception as fb_error:
                            fb_err_msg = str(fb_error)
                            if "429" in fb_err_msg or "RESOURCE_EXHAUSTED" in fb_err_msg or "503" in fb_err_msg or "UNAVAILABLE" in fb_err_msg:
                                if fb_attempt < max_retries - 1:
                                    st.warning(f"⚠️ Modello {fallback_model} sovraccarico. Attesa di {retry_delay} secondi...")
                                    time.sleep(retry_delay)
                                    continue
                            st.warning(f"❌ Errore con {fallback_model}: {fb_error}. Passo al prossimo modello...")
                            break
                
                st.error("❌ Nessun modello di backup è riuscito a completare la richiesta.")
                return []
            else:
                st.error(f"Errore durante l'identificazione degli esercizi: {e}")
                try:
                    # Fallback per rimuovere eventuali delimitatori markdown se presenti
                    text_cleaned = response.text.replace("```json", "").replace("```", "").strip()
                    data = json.loads(text_cleaned)
                    return data
                except:
                    return []
    return []

def build_system_prompt(dettaglio, anno_scolastico, indirizzo_studi=None):
    """Costruisce il system prompt personalizzato per Gemini seguendo le regole ministeriali e di stile."""
    
    # Livello di dettaglio con istruzioni e vincoli stringenti
    if "Basso" in dettaglio:
        dettaglio_istruzioni = """A. LIVELLO DI DETTAGLIO: BASSO (SOLO CALCOLI E PASSAGGI FORMALI)
- Obiettivo: Fornire una scheda di calcolo essenziale, compatta e diretta.
- VINCOLO RIGIDO (ZERO PROSA DISCORSIVA): È SEVERAMENTE VIETATO scrivere paragrafi narrativi, preamboli, frasi introduttive, spiegazioni verbali estese o commenti pedagogici.
- Struttura: Mostra ESCLUSIVAMENTE equazioni, passaggi algebrici, sostituzioni numeriche e schemi di calcolo.
- Connettivi ammessi: Usa solo formule matematiche collegate da simboli logici (es. $\\implies$, $\\iff$, $\\therefore$) e brevissime diciture schematiche essenziali (es. 'Posto $t = 5^x$:', 'Condizioni di esistenza:', 'Radici: $x_1 = -2, x_2 = -1$', 'Soluzione: $x \\le 9/4$').
- Per studi di funzione o proprietà: riporta unicamente elenchi sintetici puntati (es. 'Dominio: $D = \\mathbb{R}$', 'Immagine: $\\text{Im} = (0, +\\infty)$', 'Monotonia: strettamente crescente')."""
    elif "Medio" in dettaglio:
        dettaglio_istruzioni = """A. LIVELLO DI DETTAGLIO: MEDIO (INDICAZIONI ESSENZIALI E MOTIVAZIONE DEI PASSAGGI)
- Obiettivo: Fornire una spiegazione chiara, scorrevole e didatticamente equilibrata.
- Struttura: Riporta tutti i passaggi matematici completi accompagnati da brevi ed efficaci frasi di raccordo in italiano.
- Motivazione dei passaggi: Spiega sinteticamente il 'perché' di ogni passaggio algebrico o logico (es. 'Poiché la base dell\\'esponenziale è $0 < a < 1$, invertiamo il verso della disequazione:', 'Imponiamo che il denominatore sia diverso da zero per la definizione delle frazioni algebriche:', 'Scartiamo la soluzione negativa poiché $5^x > 0$ per ogni $x \\in \\mathbb{R}$')."""
    else:  # Alto
        dettaglio_istruzioni = """A. LIVELLO DI DETTAGLIO: ALTO (TRATTAZIONE COMPLETA E RISULTATI TEORICI CORRELATI)
- Obiettivo: Fornire una trattazione didattica e scientifica approfondita di livello avanzato.
- Inquadramento teorico obbligatorio: Oltre a tutti i calcoli dettagliati e alle motivazioni passaggio per passaggio, devi integrare l'inquadramento teorico formale dei concetti matematici e fisici coinvolti.
- Teoremi e definizioni: Enuncia esplicitamente le definizioni formali, i teoremi, le proprietà analitiche e le leggi applicate (es. classificazione algebrica della funzione; condizioni di esistenza dei radicali con indice pari vs dispari; definizione di iniettività, suriettività e invertibilità con calcolo dell'inversa; studio del segno del trinomio di secondo grado con $\\Delta$; proprietà asintotiche e limiti notevoli delle funzioni esponenziali).
- Interpretazione geometrica e critica: Descrivi il significato grafico nel piano cartesiano e analizza criticamente i risultati ottenuti."""
        
    system_prompt = f"""Sei un assistente esperto nella risoluzione di problemi di matematica e fisica.
Risolvi l'esercizio che ti viene fornito in modo impeccabile, formattando tutto il testo matematico in LaTeX standard.

REGOLE COMPORTAMENTALI:

{dettaglio_istruzioni}

B. ADATTAMENTO AL CONTESTO SCOLASTICO (Programmi ministeriali italiani):
- Anno scolastico: {anno_scolastico}
Adatta rigorosamente gli strumenti matematici, i teoremi utilizzabili, il livello di astrazione e il formalismo del linguaggio a questo profilo dello studente in conformità alle Indicazioni Nazionali del Ministero dell'Istruzione italiano. 
Ad esempio, NON utilizzare metodi avanzati (es. limiti, derivate o integrali) per classi inferiori alla quinta superiore, o per contesti in cui tali strumenti non sono previsti dal programma ministeriale, a meno che non sia strettamente necessario o richiesto in modo esplicito.

C. REGOLE SPECIFICHE PER LA FISICA E LA MATEMATICA:
1. Dichiarazione dei dati e delle costanti: All'inizio di ogni soluzione di fisica, devi creare un elenco dei dati noti estratti dal testo e delle costanti fisiche universali o tabellate adottate (es. $g = 9.81 \\text{{ m/s}}^2$), con le relative unità di misura.
2. Risoluzione simbolica e numerica (Fisica): I problemi di fisica devono essere impostati prima in forma puramente letterale/simbolica (ricavando la formula risolutiva finale espressa con le variabili) e solo successivamente in forma numerica, sostituendo i valori e calcolando il risultato.
3. Arrotondamenti e cifre significative: I risultati numerici finali devono essere arrotondati in modo coerente con i dati di partenza del problema (di norma 2 o 3 cifre significative, o in base alle preferenze indicate dall'utente).
4. Verifica aritmetica (Python sandbox): Devi utilizzare lo strumento di esecuzione del codice (Code Execution) per calcolare ed elaborare ogni espressione numerica o calcolo trigonometrico, riportando nel testo LaTeX il risultato esatto verificato per evitare allucinazioni aritmetiche.
5. Input multilingua: Se il problema originale è in lingua straniera, comprendi la traccia ma genera la spiegazione e i passaggi risolutivi esclusivamente in lingua italiana.

D. REGOLE DI FORMATTAZIONE MATEMATICA (ZERO TOLLERANZA):
1. Sintassi LaTeX standard: Ogni singola variabile, costante, numero, equazione o formula deve essere racchiusa tra delimitatori LaTeX.
2. Formule in linea: Utilizzare il singolo dollaro `$ ... $` (es. $x = 2$).
3. Equazioni isolate: Utilizzare il doppio dollaro `$$ ... $$` per equazioni importanti isolate su una riga.
4. Unità di misura: Devono essere sempre formattate in LaTeX usando il font non corsivo (es. $9.81 \\text{{ m/s}}^2$ oppure $100 \\text{{ N}}$).
5. No caratteri Unicode speciali: Non utilizzare MAI simboli matematici al di fuori di LaTeX (es. non usare 𝛂, √, 𝜋, ±, ma usa sempre `\\alpha`, `\\sqrt`, `\\pi`, `\\pm`).
6. CARATTERI RISERVATI (CRITICO): È SEVERAMENTE VIETATO usare i simboli `_`, `^`, `\\cdot` o qualsiasi altro comando matematico al di fuori dei delimitatori `$ ... $`. Ad esempio, scrivere testualmente `(c_A = 4186)` genererà un ERRORE FATALE che distrugge il documento. DEVI usare SEMPRE i dollari: `($c_A = 4186$)` oppure $c_A$. Nessuna eccezione.
7. NO PARENTESI GRAFFE VUOTE O PEDICI/APICI ORFANI: È assolutamente vietato generare apici o pedici non completati o con graffe vuote (es. non generare mai `^{{}}` o `_{{}}` o parentesi graffe aperte senza contenuto). Qualsiasi esponente o pedice deve avere sempre un valore esplicito tra le parentesi graffe (es. `^{{2}}`).

E. STRUTTURA OBBLIGATORIA DELL'OUTPUT (ZERO TOLLERANZA):
Devi imporre una struttura fissa e obbligatoria per ogni esercizio. Ogni esercizio deve essere trattato come un capitolo a sé stante. All'interno di ogni capitolo, il testo deve contenere ESCLUSIVAMENTE questi due paragrafi:

# Esercizio X
### Testo del problema
[Riporta qui la traccia originale dell'esercizio]

### Soluzione del problema
[Esposizione della risoluzione conforme al livello di dettaglio richiesto]

F. SINTESI ESTREMA:
Elimina qualsiasi preambolo generale, saluto o commento pedagogico (es. "Ecco la soluzione", "Spero sia chiaro"). La soluzione deve essere diretta, densa e concentrata unicamente sui contenuti matematici e fisici.
"""
    return system_prompt

def solve_exercise_with_gemini(api_key, system_prompt, esercizio_testo, immagini_correlate, ex_id, img_dir):
    """Risolve un singolo esercizio inviandolo a Gemini con Code Execution abilitato."""
    client = genai.Client(api_key=api_key)
    
    contents = [f"Risolvi il seguente problema (Esercizio {ex_id}):\n{esercizio_testo}"]
    
    # Aggiungi eventuali immagini correlate all'esercizio como input multimodale
    for img_name in immagini_correlate:
        path = os.path.join(img_dir, img_name)
        if os.path.exists(path):
            try:
                img = Image.open(path)
                contents.append(img)
            except Exception as e:
                st.warning(f"Impossibile aprire l'immagine {img_name} per Gemini: {e}")
                
    primary_model = "gemini-3.7-flash"
    fallback_models = ["gemini-3.5-flash", "gemini-2.5-pro", "gemini-3.1-flash-lite"]
    
    max_retries = 3
    retry_delay = 20  # secondi
    
    # 1. Tentativo con il modello primario e retry
    for attempt in range(max_retries):
        try:
            # Chiamata di risoluzione con il modello primario
            response = client.models.generate_content(
                model=primary_model,
                contents=contents,
                config=types.GenerateContentConfig(
                    system_instruction=system_prompt,
                    tools=[types.Tool(code_execution=types.ToolCodeExecution())]
                )
            )
            return response.text
        except Exception as e:
            err_msg = str(e)
            # Se la risorsa è esaurita (429) o il server è sovraccarico (503), attendi e riprova
            if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg or "503" in err_msg or "UNAVAILABLE" in err_msg:
                if attempt < max_retries - 1:
                    st.warning(f"⚠️ Modello {primary_model} sovraccarico o limite quota superato. Attesa di {retry_delay} secondi prima del tentativo {attempt + 2}/{max_retries}...")
                    time.sleep(retry_delay)
                    continue
            
            # Se fallisce per quota/sovraccarico/non trovato, andiamo in fallback
            if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg or "503" in err_msg or "UNAVAILABLE" in err_msg or "404" in err_msg:
                st.warning(f"⚠️ Quota esaurita o modello {primary_model} non disponibile per la risoluzione. Tento con i modelli di backup in cascata...")
                
                for fallback_model in fallback_models:
                    st.info(f"🔄 Tentativo con modello di backup: {fallback_model}...")
                    for fb_attempt in range(max_retries):
                        try:
                            response = client.models.generate_content(
                                model=fallback_model,
                                contents=contents,
                                config=types.GenerateContentConfig(
                                    system_instruction=system_prompt,
                                    tools=[types.Tool(code_execution=types.ToolCodeExecution())]
                                )
                            )
                            return response.text + f"\n\n*(Nota: Soluzione generata automaticamente con il modello di backup {fallback_model} a causa dei limiti di quota sul modello primario)*"
                        except Exception as fb_error:
                            fb_err_msg = str(fb_error)
                            if "429" in fb_err_msg or "RESOURCE_EXHAUSTED" in fb_err_msg or "503" in fb_err_msg or "UNAVAILABLE" in fb_err_msg:
                                if fb_attempt < max_retries - 1:
                                    st.warning(f"⚠️ Modello {fallback_model} sovraccarico. Attesa di {retry_delay} secondi...")
                                    time.sleep(retry_delay)
                                    continue
                            st.warning(f"❌ Errore con {fallback_model}: {fb_error}. Passo al prossimo modello...")
                            break
                
                return "\n*Errore: impossibile generare la soluzione anche con i modelli di backup.*\n"
            else:
                return f"\n*Errore durante la generazione della soluzione:* {e}\n"
    return "\n*Errore: impossibile generare la soluzione a causa di ripetuti limiti di quota.*\n"


# Modulo 4: Compilazione con Pandoc e UI Streamlit

import re


def sanitize_latex_for_pandoc(markdown_content):
    """
    Sanitizza Markdown+LaTeX generato da Gemini per prevenire crash fatali di pdflatex.

    Usa un approccio a due fasi:
      1. Segmenta il contenuto in blocchi testo/math analizzando i delimitatori $ e $$
      2. Applica sanitizzazione specifica per tipo a ciascun blocco

    Questo previene il problema per cui correzioni in un contesto (es. wrapping di comandi
    in $...$) rompono delimitatori in un altro contesto.
    """
    # ── Fase 1: segmentazione ──
    segments = _segment_into_math_and_text(markdown_content)

    # ── Fase 2: sanitizzazione e ricostruzione ──
    result_parts = []
    for seg_type, content in segments:
        if seg_type == 'text':
            sanitized = _sanitize_text_segment(content)
            # Previeni $$ accidentale alla giunzione di due blocchi inline math
            if result_parts and result_parts[-1].endswith('$') and sanitized.startswith('$'):
                result_parts.append(' ')
            result_parts.append(sanitized)
        elif seg_type == 'math_inline':
            sanitized = _sanitize_math_segment(content)
            if sanitized.strip():
                math_str = '$' + sanitized + '$'
                if result_parts and result_parts[-1].endswith('$'):
                    result_parts.append(' ')
                result_parts.append(math_str)
        elif seg_type == 'math_display':
            sanitized = _sanitize_math_segment(content)
            if sanitized.strip():
                result_parts.append('$$' + sanitized + '$$')

    result = ''.join(result_parts)

    # ── Fase 3: fix globali ──
    # \left/\right sbilanciati: rimuovi i prefissi per far compilare comunque
    left_count = len(re.findall(r'\\left(?![a-zA-Z])', result))
    right_count = len(re.findall(r'\\right(?![a-zA-Z])', result))
    if left_count != right_count:
        result = re.sub(r'\\left(?![a-zA-Z])', '', result)
        result = re.sub(r'\\right(?![a-zA-Z])', '', result)

    return result


# ═══════════════════════════════════════════════════════════════════
#  Segmentazione: testo ↔ math
# ═══════════════════════════════════════════════════════════════════

def _segment_into_math_and_text(content):
    """Segmenta *content* in una lista di tuple (tipo, contenuto).

    I tipi possibili sono ``'text'``, ``'math_inline'`` e ``'math_display'``.
    Il contenuto dei segmenti math **non** include i delimitatori ``$`` / ``$$``.
    """
    segments = []
    current = []
    state = 'text'       # 'text' | 'math_inline' | 'math_display'
    i = 0
    n = len(content)

    while i < n:
        if state == 'text':
            # $$ apre display math (controllato *prima* di $)
            if i + 1 < n and content[i] == '$' and content[i + 1] == '$':
                if current:
                    segments.append(('text', ''.join(current)))
                    current = []
                state = 'math_display'
                i += 2
            elif content[i] == '$':
                if current:
                    segments.append(('text', ''.join(current)))
                    current = []
                state = 'math_inline'
                i += 1
            else:
                current.append(content[i])
                i += 1

        elif state == 'math_inline':
            if content[i] == '$':
                segments.append(('math_inline', ''.join(current)))
                current = []
                state = 'text'
                i += 1
            else:
                current.append(content[i])
                i += 1

        elif state == 'math_display':
            if i + 1 < n and content[i] == '$' and content[i + 1] == '$':
                segments.append(('math_display', ''.join(current)))
                current = []
                state = 'text'
                i += 2
            else:
                current.append(content[i])
                i += 1

    # Segmento residuo
    remaining = ''.join(current)
    if state == 'text':
        if remaining:
            segments.append(('text', remaining))
    else:
        # Blocco math non chiuso: lo emettiamo comunque (verrà chiuso nella ricostruzione)
        segments.append((state, remaining))

    return segments


# ═══════════════════════════════════════════════════════════════════
#  Sanitizzazione: blocchi math
# ═══════════════════════════════════════════════════════════════════

def _sanitize_math_segment(content):
    """Sanitizza il contenuto di un blocco math (*senza* i delimitatori ``$``)."""

    # 0. $ stray all'interno di math: non dovrebbero esserci
    content = content.replace('$', '')

    # 1. \cdot dentro \text{} → crash fatale pdflatex; sostituiamo con spazio
    def _fix_text_cdot(m):
        inner = m.group(1)
        return r'\text{' + inner.replace('\\cdot', ' ') + '}'
    content = re.sub(r'\\text\{([^}]*)\}', _fix_text_cdot, content)

    # 2. Apici/pedici vuoti: ^{} e _{}
    content = re.sub(r'\^\{\s*\}', '', content)
    content = re.sub(r'_\{\s*\}', '', content)

    # 3. \sqrt vuoto o con spazzatura
    content = re.sub(r'\\sqrt(\[[^\]]*\])?\{\s*[·\u2022]*\s*\}', '', content)

    # 4. Bilanciamento graffe
    content = _balance_braces(content)

    # 5. Assicura che ogni \frac abbia esattamente 2 argomenti tra graffe
    content = _ensure_frac_args(content)

    # 6. % non escapato dentro math → commento LaTeX, crash
    content = re.sub(r'(?<!\\)%', r'\%', content)

    # 7. Bilanciamento finale di sicurezza (dopo le trasformazioni precedenti)
    content = _balance_braces(content)

    return content


def _balance_braces(content):
    """Bilancia ``{`` e ``}`` aggiungendo o rimuovendo dal fondo."""
    open_count = content.count('{')
    close_count = content.count('}')
    if open_count > close_count:
        content += '}' * (open_count - close_count)
    elif close_count > open_count:
        diff = close_count - open_count
        chars = list(content)
        removed = 0
        for j in range(len(chars) - 1, -1, -1):
            if chars[j] == '}' and removed < diff:
                chars.pop(j)
                removed += 1
                if removed == diff:
                    break
        content = ''.join(chars)
    return content


def _find_matching_brace(text, start):
    """Restituisce l'indice *dopo* la ``}`` corrispondente.

    *start* deve puntare a ``{``.  Restituisce ``-1`` se non trovata.
    """
    if start >= len(text) or text[start] != '{':
        return -1
    depth = 0
    i = start
    while i < len(text):
        if text[i] == '{':
            depth += 1
        elif text[i] == '}':
            depth -= 1
            if depth == 0:
                return i + 1
        i += 1
    return -1


def _ensure_frac_args(content):
    r"""Assicura che ogni ``\frac`` abbia esattamente 2 argomenti ``{…}{…}``.

    Se manca il secondo argomento, aggiunge ``{}``.
    Processa ricorsivamente il contenuto degli argomenti per gestire \frac annidati.
    """
    result = []
    i = 0
    n = len(content)

    while i < n:
        # Cerca \frac non seguito da altre lettere (es. non \fractal)
        if content[i:i + 5] == '\\frac' and (i + 5 >= n or not content[i + 5].isalpha()):
            result.append('\\frac')
            # NOTA: '\\frac' in Python source = stringa a 5 caratteri: \ f r a c
            i += 5

            # Salta spazi
            while i < n and content[i] == ' ':
                i += 1

            # ── primo argomento ──
            if i < n and content[i] == '{':
                end = _find_matching_brace(content, i)
                if end > 0:
                    inner = content[i + 1:end - 1]
                    result.append('{' + _ensure_frac_args(inner) + '}')
                    i = end
                else:
                    # { non chiusa → prendi il resto come primo argomento e aggiungi secondo vuoto
                    result.append(content[i:] + '}')
                    result.append('{}')
                    return ''.join(result)
            else:
                # Nessun primo argomento → aggiungi due vuoti
                result.append('{}{}')
                continue

            while i < n and content[i] == ' ':
                i += 1

            # ── secondo argomento ──
            if i < n and content[i] == '{':
                end = _find_matching_brace(content, i)
                if end > 0:
                    inner = content[i + 1:end - 1]
                    result.append('{' + _ensure_frac_args(inner) + '}')
                    i = end
                else:
                    result.append(content[i:])
                    return ''.join(result)
            else:
                # Manca secondo argomento → aggiungi vuoto
                result.append('{}')
        else:
            result.append(content[i])
            i += 1

    return ''.join(result)


# ═══════════════════════════════════════════════════════════════════
#  Sanitizzazione: segmenti testo
# ═══════════════════════════════════════════════════════════════════

def _sanitize_text_segment(content):
    """Sanitizza testo che si trova *fuori* dal math mode."""

    # 1. Escape % non protetto (commento LaTeX)
    content = re.sub(r'(?<!\\)%', r'\%', content)
    content = content.replace('\\\\%', '\\%')   # corregge doppio escape

    # 2. Variabili con underscore fuori math → $c_A$, $v_0$, ecc.
    content = re.sub(r'(?<!\$)\b([a-zA-Z]_[a-zA-Z0-9])\b(?!\$)', r'$\1$', content)

    # 3. Comandi LaTeX strutturali (\frac{…}{…}, \sqrt{…}, …) finiti in testo
    #    per colpa di delimitatori $ mancanti o sbagliati → wrappali in $…$
    content = _wrap_structural_commands_in_text(content)

    # 4. Comandi semplici standalone (\approx, \Delta, …) fuori math → $\cmd$
    #    Applichiamo solo alle parti che non sono già dentro $…$ (generate dallo step 3)
    parts = re.split(r'(\$[^$]+\$)', content)
    for idx, part in enumerate(parts):
        if not (part.startswith('$') and part.endswith('$') and len(part) > 2):
            parts[idx] = _wrap_simple_commands_in_text(part)
    content = ''.join(parts)

    return content


def _wrap_structural_commands_in_text(text):
    r"""Trova comandi LaTeX strutturali (``\frac``, ``\sqrt``, …) in testo e li avvolge in ``$…$``."""
    result = []
    i = 0
    n = len(text)

    while i < n:
        matched = False
        if text[i] == '\\' and i + 1 < n and text[i + 1].isalpha():
            # Estrai nome comando
            j = i + 1
            while j < n and text[j].isalpha():
                j += 1
            cmd = text[i + 1:j]

            if cmd == 'frac':
                end = _find_frac_end_in_text(text, i)
                if end > i:
                    result.append('$' + text[i:end] + '$')
                    i = end
                    matched = True
            elif cmd in ('sqrt', 'overline', 'underline', 'hat', 'vec',
                         'bar', 'tilde', 'mathbf', 'mathrm', 'mathit'):
                end = _find_command_with_brace_arg(text, i, j)
                if end > i:
                    result.append('$' + text[i:end] + '$')
                    i = end
                    matched = True

        if not matched:
            result.append(text[i])
            i += 1

    return ''.join(result)


def _wrap_simple_commands_in_text(text):
    r"""Avvolge comandi LaTeX standalone (``\approx``, ``\Delta``, …) trovati in testo con ``$\cmd$``."""
    simple_commands = [
        'approx', 'cdot', 'Delta', 'delta', 'pm', 'mp',
        'times', 'alpha', 'beta', 'gamma', 'pi', 'theta',
        'mu', 'sigma', 'omega', 'infty', 'neq', 'leq', 'geq',
        'sim', 'propto', 'equiv', 'nabla', 'partial',
    ]
    for cmd in simple_commands:
        # Il negative lookahead (?![a-zA-Z]) evita di catturare \le dentro \left, \ge dentro \geq, ecc.
        def _replacer(m, c=cmd):
            return '$\\' + c + '$'
            # NOTA: '$\\' in Python source = stringa '$\'  (dollar + backslash)
        text = re.sub(r'\\' + cmd + r'(?![a-zA-Z])', _replacer, text)
    return text


def _find_frac_end_in_text(text, start):
    r"""Trova la fine di ``\frac{…}{…}`` a partire da *start* (che punta a ``\``)."""
    i = start + 5          # dopo \frac
    n = len(text)
    while i < n and text[i] == ' ':
        i += 1
    if i >= n or text[i] != '{':
        return start
    end1 = _find_matching_brace(text, i)
    if end1 < 0:
        return start
    i = end1
    while i < n and text[i] == ' ':
        i += 1
    if i >= n or text[i] != '{':
        return start
    end2 = _find_matching_brace(text, i)
    if end2 < 0:
        return start
    return end2


def _find_command_with_brace_arg(text, cmd_start, cmd_name_end):
    r"""Trova la fine di ``\command[…]{…}`` (argomento opzionale ``[]`` + obbligatorio ``{}``)."""
    i = cmd_name_end
    n = len(text)
    while i < n and text[i] == ' ':
        i += 1
    # Argomento opzionale []
    if i < n and text[i] == '[':
        bracket_end = text.find(']', i)
        if bracket_end < 0:
            return cmd_start
        i = bracket_end + 1
        while i < n and text[i] == ' ':
            i += 1
    # Argomento obbligatorio {}
    if i < n and text[i] == '{':
        end = _find_matching_brace(text, i)
        if end > 0:
            return end
    return cmd_start


def generate_pandoc_output(markdown_content, output_format, reference_doc=None):
    """Compila il testo Markdown nel formato specificato (docx o pdf) usando pypandoc."""
    
    # Sanitizza il markdown per correggere errori comuni dell'AI prima della compilazione
    markdown_content = sanitize_latex_for_pandoc(markdown_content)
    
    # Create the temp markdown file in the current directory to avoid cross-drive/temp folder path issues with LaTeX
    with tempfile.NamedTemporaryFile(dir=os.getcwd(), delete=False, suffix=".md", mode="w", encoding="utf-8") as temp_md:
        temp_md.write(markdown_content)
        temp_md_path = temp_md.name
        
    output_filename = os.path.join(os.getcwd(), f"soluzioni_compilate_{int(time.time())}.{output_format}")
    
    try:
        extra_args = []
        if output_format == "docx":
            if reference_doc and os.path.exists(reference_doc):
                extra_args.append(f"--reference-doc={reference_doc}")
                
        elif output_format == "pdf":
            # Per il PDF, impostiamo margini e opzioni del motore LaTeX
            extra_args.append("--pdf-engine=pdflatex")
            extra_args.append("--pdf-engine-opt=-interaction=nonstopmode")
            extra_args.append("-V")
            extra_args.append("geometry:margin=2cm")
            
        pypandoc.convert_file(
            temp_md_path,
            output_format,
            outputfile=output_filename,
            extra_args=extra_args
        )
        
        if os.path.exists(output_filename):
            with open(output_filename, "rb") as f:
                data = f.read()
            try:
                os.remove(temp_md_path)
                os.remove(output_filename)
            except:
                pass
            return data
            
    except OSError as e:
        st.error(f"Errore di sistema (Pandoc/LaTeX non trovato): {e}\nAssicurati che Pandoc e un compilatore LaTeX (es. MiKTeX o MacTeX) siano installati sul computer e inseriti nel PATH.")
    except Exception as e:
        # Se c'è stato un errore (es. exitcode 43) ma il file PDF è stato comunque generato grazie a nonstopmode, recuperiamolo!
        if output_format == "pdf" and os.path.exists(output_filename):
            st.warning("⚠️ **Nota:** Il modello AI ha commesso alcune imprecisioni nella sintassi matematica, ma il compilatore PDF è riuscito ad auto-correggerle. Il PDF è stato generato con successo, ma potrebbe contenere lievi imperfezioni visive in alcune formule.")
            try:
                with open(output_filename, "rb") as f:
                    data = f.read()
                try:
                    os.remove(temp_md_path)
                    os.remove(output_filename)
                except:
                    pass
                return data
            except:
                pass
                
        err_msg = str(e)
        if "pdflatex" in err_msg or "exitcode \"43\"" in err_msg or "Missing $" in err_msg:
            st.error(
                "❌ **Errore fatale di formattazione LaTeX (PDF)**\n\n"
                "Il modello AI ha commesso un errore di sintassi matematica "
                "(ad esempio ha dimenticato di racchiudere una formula tra i simboli del dollaro `$`). "
                "Il motore di compilazione PDF si è bloccato e non è riuscito ad auto-recuperare.\n\n"
                "💡 **Soluzione rapida:** Seleziona **.docx (Microsoft Word)** come formato di output nella barra laterale e clicca di nuovo su 'Avvia'. "
                "Word è molto più tollerante e ti permetterà di avere subito il documento completo!\n\n"
                f"*(Dettaglio tecnico per debug: {err_msg})*"
            )
        else:
            st.error(f"Errore generico durante la compilazione Pandoc: {e}")
    finally:
        if os.path.exists(temp_md_path):
            try:
                os.remove(temp_md_path)
            except:
                pass
    return None


def run_app():
    # Inizializzazione dello stato di sessione
    if "pipeline_step" not in st.session_state:
        st.session_state.pipeline_step = "upload"
    if "review_data" not in st.session_state:
        st.session_state.review_data = None
    if "completed_data" not in st.session_state:
        st.session_state.completed_data = None

    # Caricamento automatico dell'API Key
    api_key = get_api_key()
    
    # Intestazione UI
    st.title("📐 Risolutore di Matematica & Fisica")
    
    # Barra laterale per i controlli interattivi
    with st.sidebar:
        st.header("⚙️ Configurazione")
        
        # 1. Livello di dettaglio
        dettaglio = st.selectbox(
            "Livello di dettaglio spiegazioni",
            options=["Basso (solo calcoli)", "Medio (indicazioni essenziali)", "Alto (spiegazioni ampie)"],
            index=2
        )
        
        # 2. Anno scolastico
        anno_scolastico = st.selectbox(
            "Anno scolastico",
            options=[
                "1° Anno (Prima)",
                "2° Anno (Seconda)",
                "3° Anno (Terza)",
                "4° Anno (Quarta)",
                "5° Anno (Quinta)",
                "Scuola Media (Secondaria di I grado)"
            ],
            index=4
        )
        
        # 3. Formato di output
        output_format_sel = st.selectbox(
            "Formato di output finale",
            options=[".docx (Microsoft Word)", ".pdf (Documento PDF)"],
            index=0
        )
        output_format = "docx" if ".docx" in output_format_sel else "pdf"
        
        # 4. Gestione multi-file
        multi_file_sel = st.selectbox(
            "Gestione caricamento multi-file",
            options=["Unico documento cumulativo", "File separati"],
            index=0
        )
        
        st.write("---")
        st.markdown(
            "<div style='display: flex; align-items: center; gap: 8px; color: #2D3748; font-size: 0.9rem; font-weight: 600; margin-bottom: 0.8rem;'>"
            "<span>🎓</span><span>Realizzato dal Prof. Tommaso Patti</span>"
            "</div>",
            unsafe_allow_html=True
        )
        with st.expander("🔒 Privacy & Trattamento Dati (Pilot)"):
            st.caption(
                "Questa applicazione è configurata per un pilot didattico privato tra collaboratori autorizzati:\n\n"
                "• **Nessun Salvataggio Permanente:** I file caricati vengono elaborati in memoria temporanea e distrutti alla fine della sessione.\n\n"
                "• **Anonimizzazione:** Caricare unicamente il testo dei quesiti; non includere intestazioni con dati personali, nomi o voti di studenti.\n\n"
                "• **Fornitore AI:** L'elaborazione avviene tramite API Google Gemini; nessun dato viene registrato o utilizzato per addestrare modelli pubblici."
            )
        
    # Area principale caricamento file
    st.write("Trascina o carica uno o più compiti di matematica e fisica (PDF, DOCX, TXT, PNG, JPG). L'applicazione estrarrà il testo e le immagini, identificherà i singoli quesiti e li risolverà sequenzialmente.")
    
    # Dettaglio informativo sui limiti del sistema
    st.info(
        "ℹ️ **Limiti e Guardrails del Sistema per questa sessione:**\n"
        "* 📁 **Numero file**: Massimo 5 file per volta.\n"
        "* ⚖️ **Dimensione file**: Massimo 15 MB per singolo file, 50 MB totali per caricamento.\n"
        "* 📄 **Limiti pagine**: Massimo 20 pagine per PDF.\n"
        "* 🖼️ **Risoluzione immagini**: Massimo 20 Megapixel per singola pagina o immagine.\n"
        "* ✏️ **Esercizi**: Massimo 10 problemi elaborabili per singola esecuzione."
    )
    
    uploaded_files = st.file_uploader(
        "Carica documenti o immagini degli esercizi",
        type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
        accept_multiple_files=True
    )
        
    # GESTIONE FASE 1: UPLOAD E ANALISI
    if st.session_state.pipeline_step == "upload":
        if uploaded_files:
            if not api_key:
                st.error(
                    "### ⚠️ Gemini API Key Mancante\n"
                    "Non è stata rilevata alcuna chiave API per Gemini. Per risolvere i problemi:\n\n"
                    "1. Imposta la variabile d'ambiente **`GEMINI_API_KEY`**.\n"
                    "2. Oppure configura il segreto in **`st.secrets`** (es. `.streamlit/secrets.toml`).\n"
                    "3. Ricarica la pagina del browser o riavvia l'applicazione."
                )
                st.stop()
                
            if len(uploaded_files) > 5:
                st.error("Puoi caricare un massimo di 5 file alla volta.")
                st.stop()
                
            total_size = sum(f.size for f in uploaded_files)
            if total_size > 50 * 1024 * 1024:
                st.error("La dimensione totale dei file supera il limite di 50 MB.")
                st.stop()
                
            if st.button("🔍 Analizza Documento ed Estrai Tracce", type="primary"):
                temp_work_dir = tempfile.mkdtemp(prefix="correttore_work_")
                img_dir = os.path.join(temp_work_dir, "images")
                os.makedirs(img_dir, exist_ok=True)
                
                processed_files = []
                all_extracted_images = []
                
                with st.spinner("Estrazione testi ed elementi visivi ad alta risoluzione in corso..."):
                    for uploaded_file in uploaded_files:
                        res = process_input_file(uploaded_file, temp_work_dir, img_dir)
                        processed_files.append(res)
                        all_extracted_images.extend(res["extracted_images"])
                
                combined_text = ""
                combined_images_for_gemini = []
                for pf in processed_files:
                    if pf["text"]:
                        combined_text += f"\n--- Contenuto da {pf['filename']} ---\n{pf['text']}\n"
                    combined_images_for_gemini.extend(pf["images_for_gemini"])
                    
                with st.spinner("Identificazione e separazione dei singoli esercizi con visione multimodale..."):
                    exercises = identify_exercises_with_gemini(
                        api_key, combined_text, combined_images_for_gemini, all_extracted_images
                    )
                    
                if not exercises:
                    st.error("Nessun esercizio identificato nel documento. Verifica la leggibilità del file.")
                    st.stop()
                    
                if len(exercises) > 10:
                    st.error("Sono stati identificati più di 10 esercizi. Limite superato.")
                    st.stop()
                    
                st.session_state.review_data = {
                    "exercises": exercises,
                    "img_dir": img_dir,
                    "temp_work_dir": temp_work_dir,
                    "processed_files": processed_files,
                    "all_extracted_images": all_extracted_images
                }
                st.session_state.pipeline_step = "review"
                st.rerun()

    # GESTIONE FASE 2: REVISIONE TRACCE (HUMAN-IN-THE-LOOP)
    elif st.session_state.pipeline_step == "review":
        rev_data = st.session_state.review_data
        exercises = rev_data["exercises"]
        img_dir = rev_data["img_dir"]
        
        st.success(f"✅ Rilevati con successo **{len(exercises)} problemi distinti**.")
        st.info("📝 **Revisione Tracce:** Verifica e correggi se necessario il testo o le formule matematiche estratte prima di avviare la risoluzione.")
        
        for idx, ex in enumerate(exercises):
            ex_id = ex.get("id", idx + 1)
            with st.expander(f"📌 **Esercizio {ex_id}**", expanded=True):
                # Mostra immagini associate se presenti
                for img_name in ex.get("immagini_correlate", []):
                    img_path = os.path.join(img_dir, img_name)
                    if os.path.exists(img_path):
                        st.image(img_path, caption=f"Figura associata: {img_name}", width=450)
                
                # Area testo modificabile
                current_text = st.text_area(
                    f"Testo / Formule (LaTeX) - Esercizio {ex_id}:",
                    value=ex.get("testo", ""),
                    key=f"edit_ex_{ex_id}",
                    height=130
                )
                
                st.markdown("**Anteprima Matematica Live:**")
                st.markdown(current_text)
                
        col_submit, col_cancel = st.columns([2, 1])
        with col_submit:
            if st.button("🚀 Conferma Tracce e Genera Soluzioni", type="primary"):
                # Aggiorna il testo degli esercizi con le modifiche inserite dall'utente
                for idx, ex in enumerate(exercises):
                    ex_id = ex.get("id", idx + 1)
                    if f"edit_ex_{ex_id}" in st.session_state:
                        ex["testo"] = st.session_state[f"edit_ex_{ex_id}"]
                        
                sys_prompt = build_system_prompt(dettaglio, anno_scolastico)
                
                # Risoluzione sequenziale (Modulo 2)
                solved_list = []
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for idx, ex in enumerate(exercises):
                    if idx > 0:
                        time.sleep(6.0) # Rate limit delay per evitare 429
                    ex_id = ex.get("id", idx + 1)
                    ex_text = ex.get("testo", "")
                    ex_imgs = ex.get("immagini_correlate", [])
                    
                    status_text.write(f"✍️ **Risoluzione Esercizio {ex_id} in corso...** ({idx+1}/{len(exercises)})")
                    soluzione = solve_exercise_with_gemini(api_key, sys_prompt, ex_text, ex_imgs, ex_id, img_dir)
                    
                    solved_list.append({
                        "id": ex_id,
                        "testo": ex_text,
                        "immagini": ex_imgs,
                        "soluzione": soluzione
                    })
                    progress_bar.progress((idx + 1) / len(exercises))
                    
                status_text.success("🎉 Risoluzione completata! Generazione del file finale...")
                
                # Compilazione ed esportazione dell'output (Modulo 4)
                template_docx_path = "template.docx" if os.path.exists("template.docx") else None
                
                markdown_out = "# Soluzione Esercizi di Matematica e Fisica\n\n"
                markdown_out += f"**Profilo Studente:** {anno_scolastico}\n"
                markdown_out += f"**Livello di dettaglio:** {dettaglio}\n\n"
                markdown_out += "---\n\n"
                
                for ex in solved_list:
                    soluzione_text = ex['soluzione']
                    
                    # Inserimento immagini prima della soluzione
                    immagini_md = ""
                    for img_name in ex["immagini"]:
                        img_path = os.path.join(img_dir, img_name)
                        if os.path.exists(img_path):
                            img_abs_path = os.path.abspath(img_path).replace("\\", "/")
                            immagini_md += f"![Immagine correlata]({img_abs_path})\n\n"
                            
                    if immagini_md:
                        if "### Soluzione del problema" in soluzione_text:
                            soluzione_text = soluzione_text.replace("### Soluzione del problema", immagini_md + "### Soluzione del problema")
                        else:
                            soluzione_text += "\n\n" + immagini_md
                            
                    markdown_out += f"{soluzione_text}\n\n"
                    markdown_out += "---\n\n"
                    
                if output_format == "pdf":
                    yaml_header = """---
geometry: margin=2cm
---

"""
                    markdown_out = yaml_header + markdown_out
                    
                with st.spinner(f"Compilazione del documento .{output_format} con Pandoc in corso..."):
                    file_data = generate_pandoc_output(markdown_out, output_format, template_docx_path)
                    
                st.session_state.completed_data = {
                    "file_data": file_data,
                    "output_format": output_format,
                    "file_name": f"Soluzioni_Verifica.{output_format}"
                }
                st.session_state.pipeline_step = "done"
                st.rerun()
                
        with col_cancel:
            if st.button("🔄 Annulla e Carica Nuovi File"):
                if rev_data.get("temp_work_dir") and os.path.exists(rev_data["temp_work_dir"]):
                    shutil.rmtree(rev_data["temp_work_dir"], ignore_errors=True)
                st.session_state.pipeline_step = "upload"
                st.session_state.review_data = None
                st.session_state.completed_data = None
                st.rerun()

    # GESTIONE FASE 3: DOWNLOAD ED ESPORTAZIONE
    elif st.session_state.pipeline_step == "done":
        comp_data = st.session_state.completed_data
        st.balloons()
        st.success("✅ **Il tuo documento con le soluzioni è pronto per il download!**")
        
        if comp_data and comp_data.get("file_data"):
            st.download_button(
                label=f"💾 Scarica {comp_data['file_name']}",
                data=comp_data["file_data"],
                file_name=comp_data["file_name"],
                mime="application/octet-stream",
                type="primary"
            )
            
        st.write("---")
        if st.button("🔄 Elabora un altro compito"):
            if st.session_state.review_data and st.session_state.review_data.get("temp_work_dir"):
                shutil.rmtree(st.session_state.review_data["temp_work_dir"], ignore_errors=True)
            st.session_state.pipeline_step = "upload"
            st.session_state.review_data = None
            st.session_state.completed_data = None
            st.rerun()

if __name__ == "__main__":
    run_app()
